import os
from typing import List
from langchain_core.documents import Document

class DocumentLoader:
    """
    Loads documents from files of type PDF, DOCX, Markdown, or Text.
    """
    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
        try:
            return loader.load()
        except Exception as e:
            print(f"[ERROR] Failed to load PDF file {file_path}: {e}")
            return []

    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_content = "\n".join(full_text)
            metadata = {"source": os.path.basename(file_path), "path": file_path}
            return [Document(page_content=text_content, metadata=metadata)]
        except Exception as e:
            print(f"[ERROR] Failed to load DOCX file {file_path}: {e}")
            return []

    @staticmethod
    def load_text_or_md(file_path: str) -> List[Document]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            metadata = {"source": os.path.basename(file_path), "path": file_path}
            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            print(f"[ERROR] Failed to load text/md file {file_path}: {e}")
            return []

    def load(self, file_path: str) -> List[Document]:
        """
        Detect file type and load document.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext == ".docx":
            return self.load_docx(file_path)
        elif ext in [".md", ".txt"]:
            return self.load_text_or_md(file_path)
        else:
            print(f"[WARNING] Unsupported file extension {ext} for file {file_path}. Attempting as plain text.")
            return self.load_text_or_md(file_path)
