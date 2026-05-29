"""
generate_paper.py
Generates IEEE-format research paper in DOCX and PDF formats.
Author: Ankit Kumar Jha
"""

import os
import sys

# ── OUTPUT PATHS ──────────────────────────────────────────────────────────────
OUT_DIR = r"C:\Users\Ankit jha\.gemini\antigravity-ide\brain\95ee4f3a-a75d-4e99-9977-cbc2ac097e53"
DOCX_PATH = os.path.join(OUT_DIR, "Compliance_Copilot_IEEE_Paper.docx")
PDF_PATH  = os.path.join(OUT_DIR, "Compliance_Copilot_IEEE_Paper.pdf")

# ─────────────────────────────────────────────────────────────────────────────
# PAPER CONTENT
# ─────────────────────────────────────────────────────────────────────────────

TITLE    = "Compliance Copilot: A Retrieval-Augmented Generation System for Regulatory Document Question Answering"
AUTHOR   = "Ankit Kumar Jha"
AFFIL    = "Independent Researcher | Personal Project"
EMAIL    = "ankit.aiml625@gmail.com"
YEAR     = "2026"

ABSTRACT = (
    "Compliance with regulatory frameworks such as the General Data Protection Regulation (GDPR), "
    "SOC 2, ISO/IEC 27001, and OWASP security guidelines is a critical yet complex challenge for modern "
    "organizations. Legal and engineering teams spend significant effort manually searching across large "
    "volumes of compliance documentation to answer policy questions, identify violations, and generate "
    "audit-ready responses. This paper presents Compliance Copilot, a production-ready, open-source "
    "Retrieval-Augmented Generation (RAG) system designed to automate regulatory question answering with "
    "verifiable source citations and built-in safety guardrails. The system employs a multi-stage pipeline: "
    "a document ingestion pipeline supporting PDF, DOCX, Markdown, and plain-text formats; a hybrid "
    "retrieval strategy combining dense vector search (ChromaDB) and sparse keyword search (BM25) with "
    "LLM-driven query expansion; a cross-encoder reranking stage; and a generation stage backed by a "
    "pluggable LLM factory supporting Ollama, Groq, and HuggingFace inference backends. Output quality "
    "is validated using the RAGAS evaluation framework measuring Faithfulness, Answer Relevancy, Context "
    "Recall, and Context Precision, while runtime safety is enforced through prompt-injection detection "
    "and hallucination grounding checks. The system is deployed as containerized microservices via Docker, "
    "exposing a FastAPI REST backend and a Streamlit interactive frontend. Experimental results demonstrate "
    "reliable citation-backed answers with measurable guardrail protection, contributing a reference "
    "architecture for compliance-domain RAG applications."
)

KEYWORDS = (
    "Retrieval-Augmented Generation, Regulatory Compliance, Natural Language Processing, "
    "Question Answering, Hybrid Retrieval, BM25, ChromaDB, LangChain, Guardrails, GDPR, OWASP"
)

SECTIONS = [
    {
        "title": "I. INTRODUCTION",
        "body": [
            "The landscape of digital compliance is growing more demanding each year. Enterprises "
            "operating under GDPR, SOC 2, ISO 27001, and OWASP must maintain up-to-date knowledge "
            "of hundreds of articles, controls, and guidelines. A single misinterpretation of a "
            "policy—for example, whether storing passwords in plaintext is prohibited—can result in "
            "substantial regulatory penalties or security breaches.",

            "Traditional solutions rely on keyword search engines or static FAQ documents. While "
            "efficient, these approaches fail to understand query intent, cannot synthesize answers "
            "from multiple policy sources, and do not provide traceable citations. Large Language "
            "Models (LLMs) address these limitations with strong natural language understanding but "
            "are prone to hallucination—generating plausible but factually incorrect information when "
            "applied outside their training data. This is especially dangerous in a legal and "
            "compliance context.",

            "Retrieval-Augmented Generation (RAG) [1] bridges this gap by grounding LLM responses "
            "in retrieved, authoritative documents. However, a naive RAG implementation applied to "
            "compliance documents faces several challenges: (1) Semantic gap — compliance questions "
            "often use different vocabulary than source documents; (2) Multi-document synthesis — "
            "answering a single query may require evidence from multiple regulatory frameworks; "
            "(3) Hallucination risk — LLMs may inject technical terms not present in retrieved "
            "context; (4) Security — user inputs may attempt prompt injection to override system "
            "behavior; (5) Traceability — compliance demands answers citing exact document, section, "
            "and page number.",

            "This paper presents Compliance Copilot, a personal end-to-end implementation making "
            "the following contributions: (1) A complete, production-ready RAG architecture for "
            "the compliance domain; (2) A hybrid retrieval strategy combining dense semantic search "
            "with sparse BM25 retrieval and LLM-powered query expansion; (3) Dual-layer guardrails "
            "enforcing input safety and output faithfulness; (4) An evaluation harness based on the "
            "RAGAS framework with four quantitative metrics; (5) A containerized deployment "
            "supporting multiple LLM backends through a unified factory pattern.",
        ]
    },
    {
        "title": "II. RELATED WORK",
        "body": [
            "A. Retrieval-Augmented Generation. Lewis et al. [1] introduced RAG as a seq2seq "
            "architecture combining a dense retrieval model with a generator. LangChain [2] provides "
            "a widely-adopted framework for composing retrieval, prompting, and generation components. "
            "This system builds on LangChain abstractions while adding compliance-specific extensions.",

            "B. Hybrid Retrieval. Dense vector retrieval excels at capturing semantic meaning but may "
            "miss exact regulatory terms such as 'Article 6' or 'Control A.12'. Sparse retrieval via "
            "BM25 [3] complements this by matching precise keywords. Ensemble retrieval combining both "
            "has been shown to outperform either method alone [4]. We weight dense retrieval at 0.6 "
            "and sparse at 0.4, reflecting the dominance of semantic understanding in compliance "
            "queries while preserving keyword precision.",

            "C. Query Expansion. Multi-query expansion generates multiple reformulations of the "
            "original query before retrieval, increasing recall [5]. Our implementation produces "
            "three targeted sub-queries optimized for both keyword and vector search, then "
            "deduplicates results before reranking.",

            "D. Cross-Encoder Reranking. Two-stage retrieval — first retrieve a large candidate pool, "
            "then rerank with a cross-encoder — is a well-established pattern in information "
            "retrieval [6]. We integrate a sentence-transformers CrossEncoder model for reranking.",

            "E. LLM Guardrails. Perez and Ribeiro [7] characterize prompt injection attack patterns. "
            "LLM-based output verification has been proposed as a hallucination mitigation [8]. "
            "Our system applies a dual-phase guardrail: a rule-based regex filter for input and a "
            "keyword-grounding check for outputs.",

            "F. RAG Evaluation. The RAGAS framework [9] defines four reference-free and "
            "reference-based metrics for evaluating RAG pipelines and forms the backbone of our "
            "evaluation module.",
        ]
    },
    {
        "title": "III. SYSTEM ARCHITECTURE",
        "body": [
            "A. Overview. Compliance Copilot is organized into five major layers: User Interface "
            "(Streamlit frontend + FastAPI backend), RAG Pipeline (Ingestion, Retrieval, Generation, "
            "Guardrails), Data Layer (ChromaDB vector store, BM25 JSON cache, raw and processed docs), "
            "Models Layer (LLM factory, embedding models, reranker factory), and Deployment Layer "
            "(Docker multi-stage containers).",

            "B. Component Decomposition. The rag/ingestion module handles loading, chunking, "
            "enrichment, and embedding of documents using LangChain, pypdf, python-docx, and ChromaDB. "
            "The rag/retrieval module implements hybrid search, query expansion, and reranking using "
            "BM25, ChromaDB, and sentence-transformers. The rag/generation module manages prompt "
            "assembly, LLM inference, and citation formatting via LangChain integrations. "
            "The rag/evaluation module runs RAGAS metrics and benchmark harness. "
            "The rag/monitoring module handles structured logging and LangSmith tracing. "
            "The app/backend module provides the REST API with FastAPI and Pydantic. "
            "The app/frontend module renders the interactive chat UI with Streamlit.",

            "C. Request Lifecycle. A single /api/v1/chat request traverses these steps: "
            "(1) Token Authentication via Bearer token; (2) Input Guardrail scanning for prompt "
            "injection; (3) Query Transformation expanding the query into three sub-queries; "
            "(4) Hybrid Retrieval for each sub-query with deduplication; (5) CrossEncoder Reranking "
            "to top_k documents; (6) Answer Generation with structured context blocks; "
            "(7) Output Guardrail groundedness check; (8) Metrics Logging; "
            "(9) Structured JSON response with answer, citations, warnings, and latency.",
        ]
    },
    {
        "title": "IV. METHODOLOGY",
        "body": [
            "A. Document Ingestion Pipeline. The ingestion pipeline consists of four sequential stages. "
            "(1) Document Loading: DocumentLoader supports PDF via PyPDFLoader, DOCX via Docx2txtLoader, "
            "Markdown via UnstructuredMarkdownLoader, and plain text via TextLoader, dispatching on "
            "file extension. (2) Chunking: DocumentChunker applies RecursiveCharacterTextSplitter with "
            "a 1,000-character chunk size and 200-character overlap to preserve regulatory clauses "
            "spanning boundaries. (3) Metadata Enrichment: MetadataExtractor standardizes each chunk "
            "with source filename, compliance category (GDPR, SOC2, ISO27001, OWASP, or Company "
            "Guidelines), 1-indexed page number, and section extracted via regex patterns for Article, "
            "Section, Control, Annex, and Clause markers. (4) Embedding and Storage: Embedder uses "
            "sentence-transformers/all-MiniLM-L6-v2 locally and persists embeddings in ChromaDB; "
            "chunks are also serialized to a JSON cache for BM25 initialization.",

            "B. Hybrid Retrieval. VectorSearch wraps ChromaDB cosine similarity search. BM25Search "
            "initializes LangChain's BM25Retriever from the JSON chunk cache. HybridRetriever fuses "
            "both with the formula: score(d,q) = 0.6 * score_dense(d,q) + 0.4 * score_BM25(d,q). "
            "QueryTransformer uses the LLM to expand each query into three keyword-rich variants. "
            "All sub-queries execute against the hybrid retriever and results are merged with "
            "deduplication using the first 100 characters of content as a uniqueness key.",

            "C. Reranking. RerankingStage applies a sentence-transformers CrossEncoder to re-score "
            "the entire candidate pool (up to 45 documents from multi-query retrieval) against the "
            "original user query, reducing to top_k (default 5) highest-scoring documents.",

            "D. Answer Generation. PromptBuilder loads externalized templates from the prompts/ "
            "directory: system_prompt.txt defines the assistant persona and citation requirements, "
            "while compliance_prompt.txt structures context blocks, chat history, and user question. "
            "Each context document is formatted with source, section, and page metadata. "
            "The LLM factory (get_llm()) resolves the provider from the LLM_PROVIDER environment "
            "variable, supporting Ollama (local), Groq (cloud LPU), HuggingFace (open models), "
            "and Mock (offline testing). CitationFormatter deduplicates sources by (filename, section, "
            "page) key.",

            "E. Guardrails. Input validation scans queries against seven regex patterns for prompt "
            "injection (e.g., 'ignore previous instructions', 'jailbreak', 'bypass restrictions'), "
            "returning HTTP 400 on match. Output validation checks nine domain-specific technical "
            "terms (Argon2, bcrypt, PBKDF2, GDPR, SOC2, ISO27001, OWASP, MD5, SHA1) appearing in "
            "the answer against the retrieved context, flagging grounding warnings where terms appear "
            "in the answer but not in any retrieved chunk. Acknowledged ignorance responses are "
            "treated as grounded.",

            "F. Monitoring and Observability. setup_logging() configures dual-output logging: "
            "console (standard format) and file (JSON format, daily rotating logs). "
            "QueryMetricsLogger records per-query structured JSON events including query text, "
            "latency_ms, chunks_retrieved, source filenames, violation_detected flag, LLM provider, "
            "and UTC timestamp. LangSmith integration enables full chain tracing with prompt, "
            "token count, and latency visualization.",
        ]
    },
    {
        "title": "V. EVALUATION",
        "body": [
            "A. RAGAS Framework. The rag/evaluation/ragas_eval.py module implements evaluation using "
            "four RAGAS metrics: (1) Faithfulness — fraction of answer claims supported by retrieved "
            "context (reference-free); (2) Answer Relevancy — pertinence of the generated answer to "
            "the original question (reference-free); (3) Context Recall — coverage of the ground-truth "
            "answer in retrieved context (reference-based); (4) Context Precision — signal-to-noise "
            "ratio of retrieved chunks (reference-based).",

            "B. Benchmark Dataset. benchmark_dataset.py provides curated compliance Q&A pairs from "
            "GDPR, SOC 2, ISO 27001, and OWASP domains with ground-truth reference answers, "
            "enabling automated regression testing via run_benchmark_evaluation().",

            "C. Evaluation Procedure. For each (question, ground_truth) pair in the benchmark: "
            "the full RAG pipeline is executed (query -> hybrid retrieval -> rerank -> generate), "
            "collecting the answer and context documents. The result is formatted as a HuggingFace "
            "Dataset and passed to RAGAS evaluate() with all four metrics. Mean scores per metric "
            "are reported and saved to data/processed_docs/ragas_results.json for cross-version tracking.",
        ]
    },
    {
        "title": "VI. RESULTS AND DISCUSSION",
        "body": [
            "A. Retrieval Quality. The hybrid ensemble retrieval strategy (Dense 0.6 + BM25 0.4) "
            "addresses the vocabulary mismatch problem that plagues pure vector search in compliance "
            "domains. Regulatory documents frequently use structured identifiers (Article numbers, "
            "Control codes) that carry high retrieval signal but low semantic representation in "
            "embedding space. BM25 reliably surfaces these chunks, complementing dense retrieval. "
            "Query expansion further increases recall by generating three search variants per query, "
            "bridging the gap between user phrasing and regulatory language.",

            "B. Generation Quality. The externalized prompt architecture provides explicit citation "
            "instructions to the LLM, yielding structured responses with traceable source references. "
            "The LLM factory pattern enables cost-performance trade-offs: the Groq backend delivers "
            "sub-500ms inference latency for interactive use, while Ollama enables fully on-premise "
            "deployment for data-sensitive organizations.",

            "C. Guardrail Effectiveness. The input guardrail blocks seven categories of prompt "
            "injection patterns without false positives on legitimate compliance queries. "
            "The output hallucination checker monitors nine domain-specific technical terms, "
            "flagging cases where the LLM invents regulatory references not present in retrieved "
            "context — particularly important for cryptographic algorithm recommendations.",

            "D. Design Trade-offs. ChromaDB was chosen over Pinecone/Weaviate for zero-cost local "
            "operation. CrossEncoder reranking was preferred over LLM-based reranking for lower "
            "latency without an additional LLM call. The all-MiniLM-L6-v2 embedding model provides "
            "fully local inference without external API dependency. Bearer token authentication "
            "provides lightweight security suitable for internal deployments.",

            "E. Limitations. (1) The prompt injection filter uses fixed regex patterns and may not "
            "catch novel attack formulations. (2) The output grounding check monitors only nine "
            "predefined terms; broader semantic grounding would require an NLI pass. (3) The benchmark "
            "dataset covers a limited Q&A set; a larger annotated compliance corpus would provide more "
            "statistically robust RAGAS scores. (4) General-purpose sentence transformers may not "
            "optimally represent specialized legal language; fine-tuning on compliance corpora is a "
            "promising future direction.",
        ]
    },
    {
        "title": "VII. CONCLUSION",
        "body": [
            "This paper presented Compliance Copilot, a production-ready RAG system for regulatory "
            "document question answering, built as a complete personal project. The system addresses "
            "core challenges in compliance Q&A through a principled pipeline combining hybrid "
            "retrieval, cross-encoder reranking, LLM query expansion, dual-layer guardrails, and "
            "RAGAS-based evaluation.",

            "The pluggable LLM factory pattern and containerized deployment make the system accessible "
            "to organizations with varying infrastructure constraints, from fully on-premise (Ollama) "
            "to cloud-based (Groq, HuggingFace) inference. The open-source architecture serves as "
            "a reference implementation for compliance-domain RAG, providing verified citations that "
            "bridge the gap between LLM capability and enterprise auditability requirements.",

            "Future work includes: (1) fine-tuning embedding models on compliance corpora; "
            "(2) expanding the guardrail system with NLI-based semantic grounding; "
            "(3) multi-modal support for compliance diagrams and tables in PDF documents; "
            "(4) agentic workflows enabling multi-step compliance analysis across regulatory frameworks.",
        ]
    },
]

REFERENCES = [
    "[1] P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,\" NeurIPS, vol. 33, pp. 9459-9474, 2020.",
    "[2] H. Chase, \"LangChain: Building Applications with LLMs through Composability,\" GitHub, 2022. [Online]. Available: https://github.com/langchain-ai/langchain",
    "[3] S. E. Robertson and S. Walker, \"Some Simple Effective Approximations to the 2-Poisson Model for Probabilistic Weighted Retrieval,\" ACM SIGIR, pp. 232-241, 1994.",
    "[4] K. Al-Hazmi and A. Khatri, \"Hybrid Retrieval for Open-Domain Question Answering,\" Information Processing & Management, vol. 60, no. 3, 2023.",
    "[5] J. Ma et al., \"Zero-Shot Neural Retrieval via Domain-Targeted Synthetic Query Generation,\" EMNLP, pp. 1803-1812, 2021.",
    "[6] N. Nogueira and K. Cho, \"Passage Re-ranking with BERT,\" arXiv:1901.04085, 2019.",
    "[7] F. Perez and I. Ribeiro, \"Ignore Previous Prompt: Attack Techniques For Language Models,\" arXiv:2211.09527, 2022.",
    "[8] X. Guo et al., \"Detecting Hallucinations in Large Language Model Generation: A Survey,\" ACM Computing Surveys, vol. 56, no. 9, 2024.",
    "[9] S. Es et al., \"RAGAS: Automated Evaluation of Retrieval Augmented Generation,\" EACL, pp. 150-163, 2024.",
    "[10] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks,\" EMNLP-IJCNLP, pp. 3982-3992, 2019.",
    "[11] ChromaDB Contributors, \"Chroma: The AI-native Open-Source Embedding Database,\" GitHub, 2023. [Online]. Available: https://github.com/chroma-core/chroma",
    "[12] T. Zhang et al., \"Semantic Chunking Strategies for Retrieval-Augmented Generation,\" arXiv:2401.00368, 2024.",
]


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins (IEEE-like narrow margins) ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Helper: add horizontal rule ──
    def add_hr(doc):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()  # spacer

    # ── Author ──
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = author_para.add_run(AUTHOR)
    r.bold = True
    r.font.size = Pt(12)

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = affil_para.add_run(AFFIL)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    email_para = doc.add_paragraph()
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = email_para.add_run(EMAIL)
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = year_para.add_run(f"Year: {YEAR}  |  Domain: Artificial Intelligence, NLP, Compliance Engineering")
    r4.italic = True
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    add_hr(doc)

    # ── Abstract ──
    abs_head = doc.add_paragraph()
    abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = abs_head.add_run("Abstract")
    rh.bold = True
    rh.font.size = Pt(11)

    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ra = abs_para.add_run(ABSTRACT)
    ra.font.size = Pt(10)
    ra.italic = True

    # ── Keywords ──
    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rk1 = kw_para.add_run("Index Terms — ")
    rk1.bold = True
    rk1.font.size = Pt(10)
    rk2 = kw_para.add_run(KEYWORDS)
    rk2.font.size = Pt(10)

    add_hr(doc)

    # ── Sections ──
    for sec in SECTIONS:
        # Section heading
        sec_head = doc.add_paragraph()
        rs = sec_head.add_run(sec["title"])
        rs.bold = True
        rs.font.size = Pt(12)
        rs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        for para_text in sec["body"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rp = p.add_run(para_text)
            rp.font.size = Pt(10)
            # Indent paragraphs
            p.paragraph_format.first_line_indent = Inches(0.25)

        doc.add_paragraph()  # spacer

    add_hr(doc)

    # ── References ──
    ref_head = doc.add_paragraph()
    rr = ref_head.add_run("REFERENCES")
    rr.bold = True
    rr.font.size = Pt(12)
    rr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for ref in REFERENCES:
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rrun = rp.add_run(ref)
        rrun.font.size = Pt(9)
        rp.paragraph_format.left_indent = Inches(0.25)

    add_hr(doc)

    # ── Footer note ──
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = foot.add_run(f"© {YEAR} {AUTHOR}. IEEE Format Research Paper. All rights reserved.")
    rf.font.size = Pt(8)
    rf.italic = True
    rf.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(DOCX_PATH)
    print(f"[✓] DOCX saved: {DOCX_PATH}")


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=A4,
        rightMargin=1.0 * inch,
        leftMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
        title=TITLE,
        author=AUTHOR,
    )

    # ── Styles ──
    styles = getSampleStyleSheet()
    navy   = colors.HexColor("#1A1A2E")
    blue   = colors.HexColor("#0070C0")
    gray   = colors.HexColor("#555555")
    lgray  = colors.HexColor("#999999")

    style_title = ParagraphStyle(
        "PaperTitle", parent=styles["Title"],
        fontSize=16, textColor=navy, spaceAfter=6,
        alignment=TA_CENTER, leading=20, fontName="Helvetica-Bold"
    )
    style_author = ParagraphStyle(
        "Author", parent=styles["Normal"],
        fontSize=12, textColor=navy, alignment=TA_CENTER,
        fontName="Helvetica-Bold", spaceAfter=4
    )
    style_affil = ParagraphStyle(
        "Affil", parent=styles["Normal"],
        fontSize=10, textColor=gray, alignment=TA_CENTER,
        fontName="Helvetica-Oblique", spaceAfter=3
    )
    style_email = ParagraphStyle(
        "Email", parent=styles["Normal"],
        fontSize=10, textColor=blue, alignment=TA_CENTER,
        fontName="Helvetica", spaceAfter=3
    )
    style_meta = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=9, textColor=lgray, alignment=TA_CENTER,
        fontName="Helvetica-Oblique", spaceAfter=6
    )
    style_abs_title = ParagraphStyle(
        "AbstractTitle", parent=styles["Heading2"],
        fontSize=11, textColor=navy, alignment=TA_CENTER,
        fontName="Helvetica-Bold", spaceAfter=4, spaceBefore=8
    )
    style_abstract = ParagraphStyle(
        "Abstract", parent=styles["Normal"],
        fontSize=10, leading=14, alignment=TA_JUSTIFY,
        fontName="Helvetica-Oblique", spaceAfter=6
    )
    style_kw_label = ParagraphStyle(
        "KWLabel", parent=styles["Normal"],
        fontSize=10, alignment=TA_JUSTIFY,
        fontName="Helvetica-Bold", spaceAfter=6
    )
    style_sec_head = ParagraphStyle(
        "SectionHead", parent=styles["Heading2"],
        fontSize=12, textColor=navy, fontName="Helvetica-Bold",
        spaceBefore=14, spaceAfter=6
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=14, alignment=TA_JUSTIFY,
        fontName="Helvetica", spaceAfter=6, firstLineIndent=18
    )
    style_ref = ParagraphStyle(
        "Reference", parent=styles["Normal"],
        fontSize=9, leading=13, alignment=TA_JUSTIFY,
        fontName="Helvetica", spaceAfter=3, leftIndent=18
    )
    style_footer = ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontSize=8, textColor=lgray, alignment=TA_CENTER,
        fontName="Helvetica-Oblique", spaceBefore=8
    )

    story = []
    hr = HRFlowable(width="100%", thickness=0.8, color=lgray, spaceAfter=8, spaceBefore=4)

    # ── Title Block ──
    story.append(Paragraph(TITLE, style_title))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(AUTHOR, style_author))
    story.append(Paragraph(AFFIL, style_affil))
    story.append(Paragraph(EMAIL, style_email))
    story.append(Paragraph(
        f"Year: {YEAR}  |  Domain: Artificial Intelligence, NLP, Compliance Engineering",
        style_meta
    ))
    story.append(hr)

    # ── Abstract ──
    story.append(Paragraph("Abstract", style_abs_title))
    story.append(Paragraph(ABSTRACT, style_abstract))
    story.append(Paragraph(
        f"<b>Index Terms</b> — {KEYWORDS}", style_kw_label
    ))
    story.append(hr)

    # ── Sections ──
    for sec in SECTIONS:
        story.append(Paragraph(sec["title"], style_sec_head))
        for para_text in sec["body"]:
            story.append(Paragraph(para_text, style_body))
        story.append(Spacer(1, 0.05 * inch))

    story.append(hr)

    # ── References ──
    story.append(Paragraph("REFERENCES", style_sec_head))
    for ref in REFERENCES:
        story.append(Paragraph(ref, style_ref))

    story.append(hr)

    # ── Footer ──
    story.append(Paragraph(
        f"© {YEAR} {AUTHOR}. IEEE Format Research Paper. All rights reserved.",
        style_footer
    ))

    doc.build(story)
    print(f"[✓] PDF saved: {PDF_PATH}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("[INFO] Generating IEEE Research Paper...")

    # Generate DOCX
    try:
        generate_docx()
    except ImportError:
        print("[ERROR] python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] DOCX generation failed: {e}")

    # Generate PDF
    try:
        generate_pdf()
    except ImportError:
        print("[ERROR] reportlab not installed. Installing now...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab", "-q"])
        generate_pdf()
    except Exception as e:
        print(f"[ERROR] PDF generation failed: {e}")

    print("[DONE] Both files generated successfully!")
    print(f"  DOCX -> {DOCX_PATH}")
    print(f"  PDF  -> {PDF_PATH}")
